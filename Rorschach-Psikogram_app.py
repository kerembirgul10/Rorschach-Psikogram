import streamlit as st
import pandas as pd
import gspread
from google.oauth2.service_account import Credentials
from collections import Counter
from io import BytesIO
import json
from datetime import datetime

# WORD kütüphanesi
try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    pass

# --- 1. GRUP TANIMLARI ---
GRUP_1 = ["G", "D", "Dd", "Gbl", "Dbl"]
GRUP_2 = ["F", "F+", "F-", "F+-", "FC", "FC'", "Fclob", "C", "C'", "Clob", "CF", "C'F", "ClobF", "K", "Kan", "Kob", "Kp", "E", "EF", "FE"]
GRUP_3 = ["H", "Hd", "(H)", "A", "Ad", "(A)", "Nesne", "Bitki", "Anatomi", "Coğrafya", "Doğa"]
GRUP_4 = ["Ban", "Reddetme", "Şok", "Pop", "O", "V"]
TUM_GRUPLAR = GRUP_1 + GRUP_2 + GRUP_3 + GRUP_4

# --- 2. GOOGLE SHEETS BAĞLANTISI ---
@st.cache_resource
def get_gsheet_client():
    creds_info = json.loads(st.secrets["gcp_service_account"])
    creds = Credentials.from_service_account_info(creds_info, scopes=["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"])
    return gspread.authorize(creds)

try:
    client = get_gsheet_client()
    sheet = client.open("Rorschach_Veritabani")
    user_sheet = sheet.worksheet("Kullanıcılar")
    patient_sheet = sheet.worksheet("Hastalar")
except Exception as e:
    st.error(f"Baglanti hatasi: {e}")
    st.stop()

# --- 3. TASARIM ---
st.set_page_config(page_title="Rorschach Klinik Panel", layout="wide")
st.markdown("""
    <style>
    /* Çoklu seçim kutusu kaydırma ayarı */
    .stMultiSelect div[role="listbox"] {
        max-height: 300px !important;
        overflow-y: auto !important;
        overscroll-behavior: contain !important;
    }
    textarea { border: 1px solid #ced4da !important; border-radius: 5px !important; }
    
    .metric-container {
        height: 100px; display: flex; flex-direction: column; justify-content: center; align-items: center;
        border-radius: 10px; margin-bottom: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); color: #1a1a1a;
        width: 100%;
    }
    .metric-label { font-size: 13px; font-weight: bold; margin-bottom: 2px; }
    .metric-value { font-size: 20px; font-weight: 900; }
    .c-g { background-color: #FFD93D; border: 2px solid #E2B200; }
    .c-d { background-color: #FFB347; border: 2px solid #E67E22; }
    .c-f { background-color: #FF6B6B; border: 2px solid #D63031; }
    .c-a { background-color: #A29BFE; border: 2px solid #6C5CE7; }
    .c-h { background-color: #D1A3FF; border: 2px solid #8E44AD; }
    .c-tri { background-color: #74B9FF; border: 2px solid #0984E3; }
    .c-rc { background-color: #55E6C1; border: 2px solid #20BF6B; }

    /* Alt kutu stili */
    .yanit-alt-kutu {
        padding: 15px;
        border-radius: 8px;
        border: 1px solid #eee;
        background-color: #fafafa;
        margin-bottom: 15px;
    }
    
    .footer { position: fixed; left: 0; bottom: 10px; width: 100%; text-align: center; color: #7f8c8d; font-size: 13px; }
    [data-testid="stSidebar"] { display: none; }
    button[kind="primary"] { background-color: #2ECC71 !important; color: white !important; border: none !important; }
    </style>
""", unsafe_allow_html=True)

def get_auto_height(text, min_h=68):
    if not text: return min_h
    lines = text.count('\n') + 1
    return max(min_h, lines * 28)

if 'logged_in' not in st.session_state: st.session_state['logged_in'] = False
if 'user' not in st.session_state: st.session_state['user'] = ""
if 'page' not in st.session_state: st.session_state['page'] = "Hastalarim"
if 'editing_patient' not in st.session_state: st.session_state['editing_patient'] = None
if 'form_id' not in st.session_state: st.session_state['form_id'] = datetime.now().timestamp()

# --- 4. WORD RAPOR ---
def create_word_report(h_info, calc, counts, total_r, b_cards, w_cards, b_reason, w_reason, protokol_verisi, selected_date):
    doc = Document()
    doc.add_heading(h_info['name'], 0)
    doc.add_heading('Hasta Bilgileri', level=1)
    doc.add_paragraph(f"Yaş: {h_info.get('age', '---')}\nUygulama Tarihi: {selected_date}")
    doc.add_heading('Klinik Yorumlar', level=2)
    doc.add_paragraph(h_info.get('comment', '---'))
    doc.add_heading('Kart Tercihleri', level=1)
    doc.add_paragraph(f"Beğenilen Kartlar: {str(b_cards).replace('[','').replace(']','')}\nNeden: {b_reason}")
    doc.add_paragraph(f"Beğenilmeyen Kartlar: {str(w_cards).replace('[','').replace(']','')}\nNeden: {w_reason}")
    doc.add_heading('Test Protokolü', level=1)
    table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Kart', 'Yanıt', 'Anket', 'Kodlar'
    for i, kart_yanitlari in enumerate(protokol_verisi, 1):
        for idx, y_paketi in enumerate(kart_yanitlari):
            row = table.add_row().cells
            row[0].text = f"Kart {i}" if idx == 0 else ""
            row[1].text = str(y_paketi.get('y', ''))
            row[2].text = str(y_paketi.get('a', ''))
            row[3].text = str(y_paketi.get('k', ''))
    doc.add_heading('Psikogram Analizi', level=1)
    doc.add_paragraph(f"Toplam Yanıt Sayısı (R): {total_r}")
    for k, v in calc.items(): doc.add_paragraph(f"{k}: %{v:.1f}")
    doc.add_heading('Kod Frekansları', level=2)
    for g_n, g_l in [("Lokalizasyon", GRUP_1), ("Belirleyiciler", GRUP_2), ("İçerik", GRUP_3), ("Özel Kodlar", GRUP_4)]:
        kodlar = [f"{k}: {counts[k]}" for k in g_l if counts[k] > 0]
        if kodlar: doc.add_paragraph(f"{g_n}: " + " | ".join(kodlar))
    return doc

# --- 5. ANALIZ FORMU ---
def analysis_form(edit_data=None):
    f_id = st.session_state['form_id'] if edit_data is None else f"edit_{edit_data.get('hasta_adi')}"
    st.header(f"{'Düzenle' if edit_data else 'Yeni'} Hasta Protokolü")
    
    c1, c2, c3 = st.columns([3, 1, 2])
    h_isim = c1.text_input("Hastanın Adı Soyadı", value=edit_data.get('hasta_adi', "") if edit_data else "", key=f"name_{f_id}")
    h_yas = c2.number_input("Yaş", 0, 120, value=int(edit_data.get('yas', 0)) if edit_data else 0, key=f"age_{f_id}")
    
    if edit_data and edit_data.get('tarih'):
        try: default_date = datetime.strptime(edit_data['tarih'], "%d/%m/%Y").date()
        except: default_date = datetime.now().date()
    else: default_date = datetime.now().date()
    h_tarih = c3.date_input("Test Tarihi", value=default_date, format="DD/MM/YYYY", key=f"date_{f_id}")
    tarih_str = h_tarih.strftime("%d/%m/%Y")
    h_yorum = st.text_area("Klinik Yorumlar", value=edit_data.get('klinik_yorum', "") if edit_data else "", height=100, key=f"comment_{f_id}")

    st.divider()
    def box_selector(label, key_prefix, saved_val):
        st.write(label)
        saved_list = json.loads(saved_val) if saved_val else []
        state_key = f"{key_prefix}_list_{f_id}"
        if state_key not in st.session_state: st.session_state[state_key] = saved_list
        cols = st.columns(10)
        for i in range(1, 11):
            is_sel = i in st.session_state[state_key]
            if cols[i-1].button(str(i), key=f"{state_key}_{i}", type="primary" if is_sel else "secondary"):
                if is_sel: st.session_state[state_key].remove(i)
                else: st.session_state[state_key].append(i)
                st.rerun()
        return st.session_state[state_key]

    b_cards = box_selector("En Beğendiği Kartlar", "best", edit_data.get('en_begendigi', "[]") if edit_data else "[]")
    b_reason = st.text_area("Beğenme Nedeni", value=edit_data.get('en_begendigi_neden', "") if edit_data else "", key=f"br_{f_id}")
    w_cards = box_selector("En Beğenmediği Kartlar", "worst", edit_data.get('en_beğenmediği', "[]") if edit_data else "[]")
    w_reason = st.text_area("Beğenmeme Nedeni", value=edit_data.get('en_beğenmediği_neden', "") if edit_data else "", key=f"wr_{f_id}")

    st.divider()
    
    raw_p = json.loads(edit_data['protokol_verisi']) if (edit_data and 'protokol_verisi' in edit_data) else None
    
    # KART RENKLERİ - Sadece çerçeve rengi olarak kullanılacak
    renkler = ["#2980b9", "#c0392b", "#8e44ad", "#2c3e50", "#16a085", "#27ae60", "#f39c12", "#d35400", "#c0392b", "#7f8c8d"]
    
    current_protocol = []
    cum_yanit_index = 1

    for i in range(1, 11):
        # KART ANA KUTU: Renkli Çerçeve + Başlık
        # İçerisi (background) şeffaf veya beyaz kalır, sadece border renklenir.
        st.markdown(f'''
            <div style="
                border: 3px solid {renkler[i-1]};
                border-radius: 15px;
                padding: 20px;
                margin-bottom: 30px;
                background-color: transparent;">
                <h3 style="color: {renkler[i-1]}; margin-top:0;">KART {i}</h3>
        ''', unsafe_allow_html=True)
        
        kart_key = f"kart_data_{i}_{f_id}"
        if kart_key not in st.session_state:
            if raw_p:
                try:
                    old_val = raw_p[i-1]
                    if isinstance(old_val, list): st.session_state[kart_key] = old_val
                    else: st.session_state[kart_key] = [{"y": old_val.get("yanit",""), "a": old_val.get("anket",""), "k": old_val.get("kodlar","")}]
                except: st.session_state[kart_key] = [{"y": "", "a": "", "k": ""}]
            else: st.session_state[kart_key] = [{"y": "", "a": "", "k": ""}]

        for idx, item in enumerate(st.session_state[kart_key]):
            # YANIT ALT KUTU (Gri zeminli ayrıştırma)
            st.markdown(f'<div class="yanit-alt-kutu">', unsafe_allow_html=True)
            st.write(f"**YANIT {cum_yanit_index}**")
            
            c_y, c_a = st.columns([1, 1])
            item["y"] = c_y.text_area("Yanıt Metni", value=item["y"], key=f"y_{i}_{idx}_{f_id}", height=get_auto_height(item["y"]), label_visibility="collapsed")
            item["a"] = c_a.text_area("Anket Metni", value=item["a"], key=f"a_{i}_{idx}_{f_id}", height=get_auto_height(item["a"]), label_visibility="collapsed")
            
            st.markdown("*Kod Seçimi:*")
            g_cols = st.columns(4)
            gruplar = [("Lokalizasyon", GRUP_1), ("Belirleyiciler", GRUP_2), ("İçerik", GRUP_3), ("Özel", GRUP_4)]
            
            current_codes = item["k"].split() if item["k"] else []
            selected_from_lists = []
            for g_idx, (g_name, g_list) in enumerate(gruplar):
                with g_cols[g_idx]:
                    # Multiselect her zaman görünür
                    chosen = st.multiselect(g_name, options=g_list, default=[c for c in current_codes if c in g_list], key=f"ms_{i}_{idx}_{g_idx}_{f_id}")
                    selected_from_lists.extend(chosen)

            manual_codes = [c for c in current_codes if c not in TUM_GRUPLAR]
            extra_input = st.text_input("Ekstra Kodlar", value=" ".join(manual_codes), key=f"extra_{i}_{idx}_{f_id}", placeholder="Mimari vb.")
            final_codes = selected_from_lists + extra_input.replace(",", " ").split()
            item["k"] = " ".join(list(dict.fromkeys(final_codes)))
            
            if item["k"]: st.info(f"Seçili: {item['k']}")

            if len(st.session_state[kart_key]) > 1:
                if st.button(f"Sil (Yanıt {cum_yanit_index})", key=f"del_{i}_{idx}_{f_id}"):
                    st.session_state[kart_key].pop(idx)
                    st.rerun()
            st.markdown('</div>', unsafe_allow_html=True) # Yanıt Alt Kutu Kapanış
            cum_yanit_index += 1

        if st.button(f"➕ Yanıt Ekle (Kart {i})", key=f"add_{i}_{f_id}", use_container_width=True):
            st.session_state[kart_key].append({"y": "", "a": "", "k": ""})
            st.rerun()
            
        st.markdown('</div>', unsafe_allow_html=True) # Kart Ana Kutu (Çerçeve) Kapanış
        current_protocol.append(st.session_state[kart_key])

    # Alt Menü Butonları
    c_btn1, c_btn2 = st.columns(2)
    if c_btn1.button("Sadece Kaydet", use_container_width=True):
        new_row = [st.session_state['user'], h_isim, h_yas, h_yorum, json.dumps(b_cards), json.dumps(w_cards), json.dumps(current_protocol), tarih_str, b_reason, w_reason]
        if edit_data:
            cell = patient_sheet.find(edit_data['hasta_adi']); patient_sheet.update(f'A{cell.row}:J{cell.row}', [new_row])
        else: patient_sheet.append_row(new_row)
        st.success("Kaydedildi.")

    if c_btn2.button("Psikogramı Hesapla", use_container_width=True):
        all_c = []; total_r = 0; r_8910 = 0
        for i, kart_list in enumerate(current_protocol, 1):
            for y_p in kart_list:
                if y_p["k"].strip():
                    total_r += 1
                    if i in [8, 9, 10]: r_8910 += 1
                    for c in y_p["k"].split(): all_c.append(c.strip())
        if total_r > 0:
            counts = Counter(all_c)
            calc = {"%G": (counts["G"]/total_r)*100, "%D": (counts["D"]/total_r)*100, "%F": (sum(counts[k] for k in ["F", "F+", "F-", "F+-"])/total_r)*100, "%A": ((counts["A"]+counts["Ad"])/total_r)*100, "%H": ((counts["H"]+counts["Hd"])/total_r)*100, "RC": (r_8910/total_r)*100}
            p_tri = (counts.get("FC",0)+counts.get("FC'",0)+counts.get("Fclob",0))*0.5 + (counts.get("CF",0)+counts.get("C'F",0)+counts.get("ClobF",0))*1 + (counts.get("C",0)+counts.get("C'",0)+counts.get("Clob",0))*1.5
            calc["TRI"] = (counts["K"]/p_tri)*100 if p_tri > 0 else 0
            st.subheader(f"Psikogram Analizi (R: {total_r})")
            m = st.columns(7)
            m[0].markdown(f'<div class="metric-container c-g"><div class="metric-label">%G</div><div class="metric-value">%{calc["%G"]:.0f}</div></div>', unsafe_allow_html=True)
            m[1].markdown(f'<div class="metric-container c-d"><div class="metric-label">%D</div><div class="metric-value">%{calc["%D"]:.0f}</div></div>', unsafe_allow_html=True)
            m[2].markdown(f'<div class="metric-container c-f"><div class="metric-label">%F</div><div class="metric-value">%{calc["%F"]:.0f}</div></div>', unsafe_allow_html=True)
            m[3].markdown(f'<div class="metric-container c-a"><div class="metric-label">%A</div><div class="metric-value">%{calc["%A"]:.0f}</div></div>', unsafe_allow_html=True)
            m[4].markdown(f'<div class="metric-container c-h"><div class="metric-label">%H</div><div class="metric-value">%{calc["%H"]:.0f}</div></div>', unsafe_allow_html=True)
            m[5].markdown(f'<div class="metric-container c-tri"><div class="metric-label">TRI</div><div class="metric-value">%{calc["TRI"]:.0f}</div></div>', unsafe_allow_html=True)
            m[6].markdown(f'<div class="metric-container c-rc"><div class="metric-label">RC</div><div class="metric-value">%{calc["RC"]:.0f}</div></div>', unsafe_allow_html=True)
            st.write("**Kod Frekansları:**")
            for g_n, g_l in [("Lokalizasyon", GRUP_1), ("Belirleyiciler", GRUP_2), ("İçerik", GRUP_3), ("Özel Kodlar", GRUP_4)]:
                kds = [f"{k}: {counts[k]}" for k in g_l if counts[k] > 0]
                if kds: st.write(f"**{g_n}:** " + " | ".join(kds))
            diag = create_word_report({'name':h_isim, 'age':h_yas, 'comment':h_yorum}, calc, counts, total_r, b_cards, w_cards, b_reason, w_reason, current_protocol, tarih_str)
            bio = BytesIO(); diag.save(bio)
            st.download_button("Word İndir", bio.getvalue(), f"{h_isim}.docx", use_container_width=True)

# --- 6. NAVIGASYON ---
if not st.session_state['logged_in']:
    st.title("Rorschach Klinik Panel")
    t1, t2 = st.tabs(["Giriş Yap", "Kayıt Ol"])
    with t1:
        u = st.text_input("Kullanıcı Adı", key="l_u")
        p = st.text_input("Şifre", type="password", key="l_p")
        if st.button("Giriş"):
            df = pd.DataFrame(user_sheet.get_all_records()); df.columns = df.columns.str.strip()
            if u in df['kullanici_adi'].values and str(p) == str(df[df['kullanici_adi']==u]['sifre'].values[0]):
                st.session_state['logged_in'] = True; st.session_state['user'] = u; st.rerun()
    with t2:
        nu, np, nn = st.text_input("Yeni Kullanıcı"), st.text_input("Şifre", type="password"), st.text_input("Ad Soyad")
        if st.button("Kayıt Ol"):
            user_sheet.append_row([nu, str(np), nn]); st.success("Kaydedildi.")
else:
    c_u, c_n1, c_n2, c_o = st.columns([1, 1, 1, 1])
    with c_u: st.markdown(f"#### {st.session_state['user']}")
    with c_n1:
        if st.button("Hastalarım", use_container_width=True, type="primary" if st.session_state['page']=="Hastalarim" else "secondary"):
            st.session_state['page'] = "Hastalarim"; st.session_state['editing_patient'] = None; st.rerun()
    with c_n2:
        if st.button("Yeni Hasta Ekle", use_container_width=True, type="primary" if st.session_state['page']=="Yeni Hasta Ekle" else "secondary"):
            st.session_state['page'] = "Yeni Hasta Ekle"; st.session_state['editing_patient'] = None; 
            st.session_state['form_id'] = datetime.now().timestamp(); st.rerun()
    with c_o:
        if st.button("Çıkış", use_container_width=True): st.session_state['logged_in'] = False; st.rerun()

    st.divider()

    if st.session_state['page'] == "Hastalarim":
        if not st.session_state['editing_patient']:
            search = st.text_input("", placeholder="Hasta Ara...")
            df_p = pd.DataFrame(patient_sheet.get_all_records())
            my_p = df_p[df_p['sahip'] == st.session_state['user']]
            if not my_p.empty:
                filt = my_p[my_p['hasta_adi'].str.contains(search, case=False)]
                st.markdown("---")
                h1, h2, h3 = st.columns([3, 1, 1])
                h1.write("**Hasta Adı**"); h2.write("**Rapor**"); h3.write("**İşlem**")
                for _, row in filt.iterrows():
                    r1, r2, r3 = st.columns([3, 1, 1])
                    if r1.button(row['hasta_adi'], key=f"e_{_}", use_container_width=True):
                        st.session_state['editing_patient'] = row.to_dict(); st.rerun()
                    if r2.button("📄 İndir", key=f"dl_{_}"):
                        try:
                            p_v = json.loads(row['protokol_verisi']); all_c = []; t_r = 0; r89 = 0
                            for i, kart_list in enumerate(p_v, 1):
                                if isinstance(kart_list, dict): kart_list = [kart_list]
                                for y_p in kart_list:
                                    kd = y_p.get('k','').strip() if 'k' in y_p else y_p.get('kodlar','').strip()
                                    if kd:
                                        t_r += 1
                                        if i in [8,9,10]: r89 += 1
                                        for c in kd.split(): all_c.append(c.strip())
                            counts = Counter(all_c)
                            doc = create_word_report({'name':row['hasta_adi'],'age':row['yas'],'comment':row['klinik_yorum']}, {}, counts, t_r, row['en_begendigi'], row['en_beğenmediği'], row['en_begendigi_neden'], row['en_beğenmediği_neden'], p_v, row['tarih'])
                            bio = BytesIO(); doc.save(bio)
                            st.download_button("Dosyayı Al", bio.getvalue(), f"{row['hasta_adi']}.docx")
                        except: r2.write("Hata")
                    if r3.button("🗑️ Sil", key=f"d_{_}"):
                        cell = patient_sheet.find(row['hasta_adi']); patient_sheet.delete_rows(cell.row); st.rerun()
            else: st.info("Kayıt yok.")
        else:
            if st.button("← Geri Dön", type="primary"): st.session_state['editing_patient'] = None; st.rerun()
            analysis_form(st.session_state['editing_patient'])
            
    elif st.session_state['page'] == "Yeni Hasta Ekle": 
        analysis_form()

st.markdown('<div class="footer">Kerem Birgül</div>', unsafe_allow_html=True)
